import google.generativeai as genai
from django.conf import settings
import os

# ==================================================
# Configure Gemini
# ==================================================
genai.configure(api_key=settings.GEMINI_API_KEY)

# Choose model
model = genai.GenerativeModel("gemini-2.5-flash")


# ==================================================
# Generate Assignment
# ==================================================
def generate_assignment(
    board, class_name, chapter, num_questions, subject="", difficulty="Medium"
):
    prompt = f"""
Create a school assignment strictly based on the following details:

Board: {board}
Class: {class_name}
Chapter: {chapter}
Number of questions: {num_questions}
Difficulty Level: {difficulty}
Subject: {subject}


STRICT FORMATTING RULES — follow exactly:
1. Write ALL math expressions using LaTeX notation:
   - Fractions: \\frac{{numerator}}{{denominator}}  → example: \\frac{{3}}{{4}}
   - Square root: \\sqrt{{x}}
   - Powers: x^{{2}}
   - Inline math: wrap in single dollar signs like $\\frac{{1}}{{2}}$
   - Display math (standalone equation): wrap in double dollar signs like $$x^2 + y^2 = z^2$$
2. Write questions exactly like a printed NCERT textbook — clear, simple English sentences.
3. Do NOT write raw fractions like 3/4 or (3)/(4) — always use $\\frac{{3}}{{4}}$
4. Do NOT write "sqrt(x)" — always use $\\sqrt{{x}}$
5. Start directly from Question 1. No preamble, no title, no extra text.
6. Number each question: 1.  2.  3.  etc.
7. Sub-parts should be labeled: (i), (ii), (iii) etc.
8. Keep each question on its own line with a blank line between questions.
9. Follow NCERT style — practical, real-world word problems where appropriate.
10. Do NOT add answers or solutions.
11. Difficulty must match: Easy = direct formula-based, Medium = multi-step problems, Hard = proof/application/higher-order thinking.
12. Questions must be appropriate for {subject} subject, Class {class_name}.
 
Example of CORRECT format:
1.  A fraction becomes $\\frac{{9}}{{11}}$ when 2 is added to both the numerator and denominator. Find the original fraction.
 
2.  Solve for $x$:
$$\\frac{{1}}{{x+4}} - \\frac{{1}}{{x-7}} = \\frac{{11}}{{30}}, \\quad x \\neq -4, 7$$
 
Now generate {num_questions} questions for Chapter: {chapter}, Class {class_name}.

"""

    response = model.generate_content(prompt)
    return response.text.strip()


# ==================================================
# Extract Text From File → Supports All File Types
# ==================================================
def extract_text_from_file(uploaded_file_path):
    """
    Extracts text from:
    - PDF        → Gemini
    - DOCX       → python-docx (local)
    - DOC        → python-docx (local, best effort)
    - TXT        → direct read (local)
    - XLSX       → openpyxl (local)
    - JPEG / JPG → Gemini
    - PNG        → Gemini
    """

    if not os.path.exists(uploaded_file_path):
        return ""

    file_lower = uploaded_file_path.lower()

    # --------------------------------------------------
    # TXT — direct read,
    # --------------------------------------------------
    if file_lower.endswith(".txt"):
        try:
            with open(uploaded_file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""

    # --------------------------------------------------
    # DOCX — extract using python-docx

    # --------------------------------------------------
    if file_lower.endswith(".docx"):
        try:
            import docx

            doc = docx.Document(uploaded_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception:
            return ""

    # --------------------------------------------------
    # DOC — extract using python-docx (best effort)

    # --------------------------------------------------
    if file_lower.endswith(".doc"):
        try:
            import docx

            doc = docx.Document(uploaded_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception:
            return ""

    # --------------------------------------------------
    # XLSX — extract using openpyxl

    # --------------------------------------------------
    if file_lower.endswith(".xlsx"):
        try:
            import openpyxl

            wb = openpyxl.load_workbook(uploaded_file_path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = "  ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        lines.append(row_text)
            return "\n".join(lines)
        except Exception:
            return ""

    # --------------------------------------------------
    # PDF, PNG, JPEG, JPG → Gemini Vision
    # --------------------------------------------------
    if file_lower.endswith(".pdf"):
        mime_type = "application/pdf"
    elif file_lower.endswith(".png"):
        mime_type = "image/png"
    elif file_lower.endswith(".jpg") or file_lower.endswith(".jpeg"):
        mime_type = "image/jpeg"
    else:
        # Unsupported file type — return empty safely
        return ""

    prompt_extract = """
Extract all readable handwritten or printed text from this file.
Return ONLY the text.
Do NOT explain anything.
"""

    try:
        with open(uploaded_file_path, "rb") as f:
            response = model.generate_content(
                [
                    prompt_extract,
                    {
                        "mime_type": mime_type,
                        "data": f.read(),
                    },
                ]
            )
        return response.text.strip()
    except Exception:
        return ""


# ==================================================
# Evaluate Submission → 1 API CALL
# ==================================================
def evaluate_submission(assignment_text, student_answer, max_marks=10):
    """
    Evaluates student answer question-by-question.
    Always returns marks + feedback.
    """

    prompt_evaluate = f"""
You are a strict school teacher evaluating a student's assignment.

Total Marks: {max_marks}

MARKING RULES:
- First count the total number of questions in the assignment below.
- Each question carries EQUAL marks = {max_marks} divided by total number of questions.
- If a question is fully correct, give full marks for that question.
- If a question is partially correct, give partial marks for that question.
- If a question is NOT attempted at all, give 0 for that question.
- Add up all question marks to get the final total.
- Final marks MUST NOT exceed {max_marks}.
- Final format must be exactly: Marks: <number>/{max_marks}

Assignment:
{assignment_text}

Student Answer:
{student_answer}

Return EXACT format:

Detailed Analysis:
<Question-wise feedback, show marks awarded per question like "Q1: 8/10 — reason">

Overall Feedback:
<Summary>

Marks: <number>/{max_marks}
"""

    result = model.generate_content(prompt_evaluate)
    return result.text.strip()


# ==================================================
# Main Processing Function
# ==================================================
def process_and_evaluate_submission(
    assignment_text,
    student_answer_text=None,
    uploaded_file_path=None,
    max_marks=10,
):
    """
    Supported file types:
    - PDF, DOCX, DOC, TXT, XLSX → text extraction
    - JPEG, JPG, PNG             → Gemini Vision OCR

    Maximum API Calls:
    - 0 calls → If no answer
    - 1 call  → If only text answer
    - 1 call  → If only image/PDF file (OCR + eval combined in 2 calls)
    - 2 calls → File (OCR) + Evaluation
    """

    extracted_text = ""

    # ==========================
    # Step 1: Extract from file
    # ==========================
    if uploaded_file_path:
        extracted_text = extract_text_from_file(uploaded_file_path)

    # ==========================
    # Combine Answers
    # ==========================
    final_student_answer = ""

    if student_answer_text:
        final_student_answer += student_answer_text.strip()

    if extracted_text:
        final_student_answer += "\n" + extracted_text

    # ==========================
    # If no answer → Skip AI
    # ==========================
    if not final_student_answer.strip():
        return {
            "evaluation": "Marks: 0\n\nFeedback:\nNo answer submitted.",
            "extracted_text": extracted_text,
            "final_answer": final_student_answer,
        }

    # ==========================
    # Step 2: Evaluate
    # ==========================
    evaluation_result = evaluate_submission(
        assignment_text,
        final_student_answer,
        max_marks,
    )

    return {
        "evaluation": evaluation_result,
        "extracted_text": extracted_text,
        "final_answer": final_student_answer,
    }
